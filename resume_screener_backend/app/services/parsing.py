import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []

    # wrap file_bytes in io.BytesIO() so pdfplumber can treat it like a file
    with io.BytesIO(file_bytes) as pdf_stream:
        with pdfplumber.open(pdf_stream) as pdf:
            # loop through pdf.pages, call .extract_text() on each
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:  # some pages (e.g. scanned images) return None
                    text_parts.append(page_text)

    # join the results together
    full_text = "\n".join(text_parts).strip()

    # if the final text is empty/blank, raise a clear error
    if not full_text:
        raise ValueError(
            "No extractable text found in this PDF. It may be a scanned "
            "image rather than real text."
        )

    return full_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    text_parts = []

    # wrap file_bytes in io.BytesIO() so Document() can treat it like a file
    with io.BytesIO(file_bytes) as docx_stream:
        doc = Document(docx_stream)
        # loop through doc.paragraphs, grab .text from each
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

    # join them together
    full_text = "\n".join(text_parts).strip()

    # if empty, raise a clear error
    if not full_text:
        raise ValueError(
            "No extractable text found in this DOCX file. It may be "
            "empty or contain only images/tables."
        )

    return full_text